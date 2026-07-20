import os
import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_agnostic_template(output_path):
    """
    Zero-Else compatible implementation.
    """
    if not output_path:
        return {"status": "error", "message": "output_path is empty"}

    try:
        doc = Document()
        
        # Title and Subtitle
        title = doc.add_heading("{{ titulo }}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph("{{ subtitulo }}")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(14)
        subtitle.runs[0].font.bold = True
        
        doc.add_paragraph()
        
        # Meta Info
        doc.add_paragraph("Fecha: {{ fecha }}")
        doc.add_paragraph("Tipo de Documento: {{ tipo_documento }}")
        doc.add_paragraph("Institución: {{ institucion }}")
        
        doc.add_paragraph()
        
        # Team
        doc.add_heading("Integrantes del Proyecto", level=1)
        
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        
        # Jinja2 logic for table rows
        row1 = table.rows[0]
        row1.cells[0].text = "{% tr for i in integrantes %}"
        row1.cells[1].text = ""
        
        row2 = table.rows[1]
        row2.cells[0].text = "{{ i.nombre_completo }}"
        row2.cells[1].text = "{{ i.rol }}{% tr endfor %}"
        
        doc.add_paragraph()
        
        # Main Content Injection Point
        doc.add_paragraph("{{ contenido }}")
        
        # Save template
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        
        return {"status": "success", "message": f"Template generated at {output_path}"}
    except Exception as e:
        return {"status": "error", "message": str(e)}

if __name__ == "__main__":
    # Ensure it writes to the correct root folder
    # this file is at: .qdd/core/plugins/qdd_docs_plugin/main.py
    # root is 4 levels up
    current_dir = os.path.dirname(os.path.abspath(__file__))
    root_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(current_dir))))
    output_file = os.path.join(root_dir, "plantilla", "Formato_Agnostico_QDD.docx")
    
    result = create_agnostic_template(output_file)
    print(json.dumps(result))
