import os
import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_agnostic_template(output_path):
    """
    Zero-Else compatible implementation.
    """
    if not output_path:
        return {"status": "error", "message": "output_path is empty"}

    try:
        doc = Document()
        
        # Title and Subtitle
        title = doc.add_heading("{{ titulo }}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph("{{ subtitulo }}")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(14)
        subtitle.runs[0].font.bold = True
        
        doc.add_paragraph()
        
        # Meta Info
        doc.add_paragraph("Fecha: {{ fecha }}")
        doc.add_paragraph("Tipo de Documento: {{ tipo_documento }}")
        doc.add_paragraph("Institución: {{ institucion }}")
        
        doc.add_paragraph()
        
        # Team
        doc.add_heading("Integrantes del Proyecto", level=1)
        
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        
        # Jinja2 logic for table rows
        row1 = table.rows[0]
        row1.cells[0].text = "{% for i in integrantes %}"
        row1.cells[1].text = ""
        
        row2 = table.rows[1]
        row2.cells[0].text = "{{ i.nombre_completo }}"
        row2.cells[1].text = "{{ i.rol }}{% endfor %}"
        
        doc.add_paragraph()
        
        # Main Content Injection Point
        doc.add_heading("Resumen Ejecutivo", level=1)
        doc.add_paragraph("{{ contenido }}")
        
        # Database
        doc.add_heading("Diccionario de Datos", level=1)
        doc.add_paragraph("{% for table in database %}")
        doc.add_heading("Tabla: {{ table.name }}", level=2)
        doc.add_paragraph("Descripción: {{ table.description }}")
        
        db_table = doc.add_table(rows=2, cols=3)
        db_table.style = 'Table Grid'
        hdr_cells = db_table.rows[0].cells
        hdr_cells[0].text = "Columna"
        hdr_cells[1].text = "Tipo"
        hdr_cells[2].text = "Descripción"
        
        db_row1 = db_table.rows[1]
        db_row1.cells[0].text = "{% for col in table.columns %}"
        db_row1.cells[1].text = ""
        db_row1.cells[2].text = ""
        
        db_row2 = db_table.add_row()
        db_row2.cells[0].text = "{{ col.name }}"
        db_row2.cells[1].text = "{{ col.type }}"
        db_row2.cells[2].text = "{{ col.description }}{% endfor %}"
        
        doc.add_paragraph("{% endfor %}")
        
        # Endpoints
        doc.add_heading("Catálogo de Endpoints", level=1)
        doc.add_paragraph("{% for ep in endpoints %}")
        doc.add_heading("Endpoint: {{ ep.method }} {{ ep.path }}", level=2)
        doc.add_paragraph("Descripción: {{ ep.description }}")
        doc.add_paragraph("{% endfor %}")
        
        # Save template
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        
        return {"status": "success", "message": f"Template generated at {output_path}"}
    except Exception as e:
        return {"status": "error", "message": str(e)}

if __name__ == "__main__":
    current_dir = os.path.dirname(os.path.abspath(__file__))
    root_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(current_dir))))
    output_file = os.path.join(root_dir, "plantilla", "Formato_Agnostico_QDD.docx")
    result = create_agnostic_template(output_file)
    print(json.dumps(result))
